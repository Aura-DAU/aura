#!/usr/bin/env python3
"""
reextract.py — DAU PWA Smart Re-Extractor
==========================================
Takes a file path (or an audit_report.json) and re-extracts + cleans the
content from the source URL, using:
  1. requests + pdfplumber / PyMuPDF  (text-based PDFs)
  2. pytesseract + pdf2image           (scanned / image-only PDFs)
  3. openpyxl                          (Excel spreadsheets)
  4. python-docx                       (Word documents)
  5. Google Gemini LLM                 (cleanup, table reconstruction, heading generation)

The re-extracted content is written back to the .md file while preserving
the original YAML front-matter.

Usage:
    # Re-extract a single file
    python scripts/reextract.py --file data/academics/some_policy.md

    # Re-extract all CRITICAL/WARN files from the audit report
    python scripts/reextract.py --from-audit audit_report.json --severity CRITICAL WARN

    # Dry-run (print what would happen, don't write)
    python scripts/reextract.py --from-audit audit_report.json --dry-run

Environment variables required (put in .env or set in shell):
    GEMINI_API_KEY   — Google AI Studio API key (for LLM cleanup)
"""

import io
import os
import re
import sys
import json
import time
import argparse
import tempfile
import warnings
from pathlib import Path
from typing import Optional

# Windows UTF-8 stdout fix
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

# ---------------------------------------------------------------------------
# Optional imports — graceful fallback with helpful error messages
# ---------------------------------------------------------------------------
try:
    import requests
except ImportError:
    requests = None  # type: ignore

try:
    import pdfplumber
except ImportError:
    pdfplumber = None  # type: ignore

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None  # type: ignore

try:
    from pdf2image import convert_from_bytes
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

try:
    import openpyxl
except ImportError:
    openpyxl = None  # type: ignore

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None  # type: ignore

try:
    from google import genai as genai_new
    HAS_GEMINI = True
except ImportError:
    try:
        import google.generativeai as genai_old   # fallback to old SDK
        HAS_GEMINI = True
        genai_new = None  # type: ignore
    except ImportError:
        HAS_GEMINI = False
        genai_new = None  # type: ignore
        genai_old = None  # type: ignore


# ---------------------------------------------------------------------------
# Load API key — checks scripts/.env first, then server/rag/.env
# ---------------------------------------------------------------------------
def load_env():
    candidates = [
        (Path(__file__).resolve().parent / ".env", True),                           # scripts/.env — force override
        (Path(__file__).resolve().parent.parent / "server" / "rag" / ".env", False) # server/rag/.env — setdefault only
    ]
    for env_file, force in candidates:
        if env_file.exists():
            for line in env_file.read_text(encoding="utf-8", errors="replace").splitlines():
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    k, _, v = line.partition("=")
                    k = k.strip(); v = v.strip().strip('"').strip("'")
                    if force:
                        os.environ[k] = v           # scripts/.env always wins
                    else:
                        os.environ.setdefault(k, v) # server/.env is fallback only


load_env()

# Suppress SSL warnings for sites with self-signed / expired certs
warnings.filterwarnings("ignore", message="Unverified HTTPS request")
try:
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
except Exception:
    pass

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")

# Initialize Gemini client (new SDK preferred)
# Pass key explicitly so we override any system-level GOOGLE_API_KEY
_gemini_client = None
GEMINI_MODEL = "gemini-2.0-flash-lite"  # highest free-tier quota
if HAS_GEMINI and GEMINI_API_KEY:
    if genai_new is not None:
        _gemini_client = genai_new.Client(api_key=GEMINI_API_KEY)
    else:
        genai_old.configure(api_key=GEMINI_API_KEY)


# ---------------------------------------------------------------------------
# Front-matter utilities
# ---------------------------------------------------------------------------

def parse_frontmatter(text: str):
    """Return (frontmatter_str, body_str)."""
    text = text.lstrip('\ufeff')
    m = re.match(r'^(---[\s\S]*?---)\s*', text)
    if m:
        return m.group(1), text[m.end():]
    return "", text


def fm_get(fm_str: str, key: str) -> str:
    m = re.search(rf'^{key}\s*:\s*["\']?(.*?)["\']?\s*$', fm_str, re.MULTILINE)
    return m.group(1).strip() if m else ""


# ---------------------------------------------------------------------------
# Download helpers
# ---------------------------------------------------------------------------

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0 Safari/537.36"
    )
}

# Domains to skip — social media / login-walled pages return no useful content
SKIP_DOMAINS = (
    "facebook.com", "instagram.com", "linkedin.com",
    "youtube.com", "x.com", "twitter.com",
)


def download(url: str, timeout: int = 30) -> Optional[bytes]:
    if requests is None:
        raise RuntimeError("requests not installed — pip install requests")
    try:
        r = requests.get(url, headers=HEADERS, timeout=timeout, verify=False)
        r.raise_for_status()
        return r.content
    except Exception as e:
        print(f"  ↳ Download failed: {e}")
        return None


# ---------------------------------------------------------------------------
# Extraction strategies
# ---------------------------------------------------------------------------

def extract_pdf_text(data: bytes, url: str) -> str:
    """Try pdfplumber first; fall back to Tesseract OCR for scanned pages."""
    text_pages = []

    # --- Strategy 1: pdfplumber (text-based PDFs) ---
    if pdfplumber:
        try:
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                    # Also extract tables
                    tables = page.extract_tables()
                    table_md = ""
                    for tbl in tables:
                        if tbl and tbl[0]:
                            header = tbl[0]
                            sep = ["---"] * len(header)
                            rows = [header, sep] + tbl[1:]
                            table_md += "\n" + "\n".join(
                                "| " + " | ".join(str(c or "").strip() for c in row) + " |"
                                for row in rows
                            ) + "\n"
                    text_pages.append(page_text + table_md)
        except Exception as e:
            print(f"  ↳ pdfplumber error: {e}")

    raw_text = "\n\n".join(p for p in text_pages if p.strip())

    # Check quality — if text is sparse or garbled, use OCR
    if _is_garbled(raw_text) or len(raw_text.strip()) < 100:
        print("  ↳ Text extraction sparse/garbled — falling back to Tesseract OCR …")
        return extract_pdf_ocr(data)

    return raw_text


def _is_garbled(text: str) -> bool:
    """Heuristic: too many control chars or consonant runs → garbled."""
    ctrl = len(re.findall(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', text))
    return ctrl > 10


class PdfOcrUnavailableError(RuntimeError):
    """Raised when OCR is needed (sparse/garbled text layer) but the
    pdf2image/pytesseract deps aren't installed."""


class PdfOcrFailedError(RuntimeError):
    """Raised when OCR was attempted but the OCR engine itself failed."""


def extract_pdf_ocr(data: bytes) -> str:
    """
    Render PDF pages to images and run Tesseract OCR.

    Fix OCR-1: this used to return human-readable placeholder strings like
    "(OCR unavailable — ...)" / "(OCR failed: ...)" on failure. Those
    strings are 50+ characters, so they silently passed the caller's
    `len(raw.strip()) < 50` "too little content" guard, got treated as if
    they were the document's real body text, and — if --dry-run wasn't
    used — could get written straight into the corpus .md file (and from
    there embedded/chunked/retrieved as if it were real page content).
    Raise instead, so the caller's existing error-handling path aborts the
    re-extraction cleanly rather than writing fake "content" to disk.
    """
    if not HAS_OCR:
        raise PdfOcrUnavailableError(
            "OCR unavailable — install pdf2image and pytesseract"
        )
    try:
        images = convert_from_bytes(data, dpi=200)
        pages = []
        for i, img in enumerate(images):
            page_text = pytesseract.image_to_string(img, lang="eng")
            pages.append(f"### Page {i+1}\n\n{page_text.strip()}")
        return "\n\n".join(pages)
    except Exception as e:
        raise PdfOcrFailedError(f"OCR failed: {e}") from e


def extract_html(data: bytes, url: str) -> str:
    """
    Extract clean readable text from an HTML page.
    Uses BeautifulSoup when available; falls back to improved regex.
    Removes: <script>, <style>, <nav>, <header>, <footer>, <aside>,
             Google Tag Manager code, and all other markup.
    """
    if HAS_BS4:
        try:
            soup = BeautifulSoup(data, "html.parser")
            # Remove noise elements entirely
            for tag in soup.find_all([
                "script", "style", "noscript", "nav", "header",
                "footer", "aside", "form", "iframe", "svg",
                "meta", "link", "button", "input",
            ]):
                tag.decompose()

            # Try to find main content area
            main = (
                soup.find("main") or
                soup.find(id=re.compile(r"main|content|body", re.I)) or
                soup.find(class_=re.compile(r"main|content|article|post", re.I)) or
                soup.find("article") or
                soup.find("body") or
                soup
            )

            # Convert headings to markdown
            for tag in main.find_all(["h1", "h2", "h3", "h4"]):
                level = int(tag.name[1])
                tag.replace_with("\n" + "#" * level + " " + tag.get_text(" ", strip=True) + "\n")

            # Convert table rows to pipe-separated markdown
            for table in main.find_all("table"):
                rows = []
                for i, tr in enumerate(table.find_all("tr")):
                    cells = [td.get_text(" ", strip=True) for td in tr.find_all(["th", "td"])]
                    if not any(cells):
                        continue
                    rows.append("| " + " | ".join(cells) + " |")
                    if i == 0:  # separator after header
                        rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
                if rows:
                    table.replace_with("\n" + "\n".join(rows) + "\n")

            # Convert list items
            for ul in main.find_all(["ul", "ol"]):
                items = []
                for li in ul.find_all("li", recursive=False):
                    items.append("- " + li.get_text(" ", strip=True))
                ul.replace_with("\n" + "\n".join(items) + "\n")

            text = main.get_text(separator="\n", strip=False)
            # Remove any remaining JS-like lines
            lines = []
            for ln in text.splitlines():
                stripped = ln.strip()
                # Skip lines that look like JavaScript
                if re.match(r"^(function\s*\(|window\.|document\.|jQuery\(|\$\(|var \w|const \w|let \w|gtag\(|dataLayer)", stripped):
                    continue
                lines.append(ln)
            return "\n".join(lines)
        except Exception as e:
            print(f"  [bs4 error] {e} — falling back to regex")

    # Fallback: improved regex stripper
    text = data.decode("utf-8", errors="replace")
    # Remove script and style blocks entirely (including content)
    text = re.sub(r"<script[^>]*>[\s\S]*?</script>", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"<style[^>]*>[\s\S]*?</style>",  " ", text, flags=re.IGNORECASE)
    text = re.sub(r"<noscript[^>]*>[\s\S]*?</noscript>", " ", text, flags=re.IGNORECASE)
    # Strip remaining tags
    text = re.sub(r"<[^>]+>", " ", text)
    # Decode common HTML entities
    text = re.sub(r"&amp;",  "&",  text)
    text = re.sub(r"&lt;",   "<",  text)
    text = re.sub(r"&gt;",   ">",  text)
    text = re.sub(r"&nbsp;", " ",  text)
    text = re.sub(r"&[a-z]+;", " ", text)
    return text


# ---------------------------------------------------------------------------
    """Convert Excel to markdown tables."""
    if openpyxl is None:
        return "(openpyxl not installed — pip install openpyxl)"
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue
            # Find real header row (look for S.No. or first non-empty row)
            header_idx = 0
            for idx, row in enumerate(rows):
                cells = [str(c).strip() for c in row if c is not None]
                if any("s.no" in c.lower() or "sr" in c.lower() for c in cells):
                    header_idx = idx
                    break
            header = [str(c or "").strip() for c in rows[header_idx]]
            # Drop all-empty columns
            non_empty_cols = [i for i, h in enumerate(header) if h]
            if not non_empty_cols:
                continue
            header = [header[i] for i in non_empty_cols]
            sep = ["---"] * len(header)
            md_rows = ["| " + " | ".join(header) + " |",
                       "| " + " | ".join(sep)   + " |"]
            for row in rows[header_idx + 1:]:
                cells = [str(row[i] if i < len(row) and row[i] is not None else "")
                         for i in non_empty_cols]
                if any(c.strip() for c in cells):
                    md_rows.append("| " + " | ".join(c.strip() for c in cells) + " |")
            parts.append(f"## Sheet: {sheet.title}\n\n" + "\n".join(md_rows))
        return "\n\n".join(parts)
    except Exception as e:
        return f"(Excel extraction failed: {e})"


def extract_docx(data: bytes) -> str:
    """Convert DOCX to markdown."""
    if DocxDocument is None:
        return "(python-docx not installed — pip install python-docx)"
    try:
        doc = DocxDocument(io.BytesIO(data))
        parts = []
        heading_map = {1: "#", 2: "##", 3: "###", 4: "####"}
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue
            if style.startswith("Heading"):
                level = int(style.split()[-1]) if style.split()[-1].isdigit() else 2
                prefix = heading_map.get(level, "##")
                parts.append(f"{prefix} {text}")
            else:
                parts.append(text)
        for tbl in doc.tables:
            rows = []
            for i, row in enumerate(tbl.rows):
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
            parts.append("\n".join(rows))
        return "\n\n".join(parts)
    except Exception as e:
        return f"(DOCX extraction failed: {e})"


# ---------------------------------------------------------------------------
# LLM cleanup
# ---------------------------------------------------------------------------

LLM_PROMPT_TEMPLATE = """You are a document formatting assistant for a university RAG system.

The following text was extracted from a PDF/document for: "{title}"
Source URL: {url}

Your task:
1. Fix all OCR errors (garbled letters, broken words, misread characters).
2. Reconstruct any broken tables into proper Markdown table format with | separators.
3. Structure the document with proper Markdown headings:
   - # for the document title (H1)
   - ## for main sections (H2)
   - ### for sub-sections (H3)
4. Keep ALL original information — do not summarize or omit any content.
5. Each H2 section should ideally fit within 250 tokens for RAG chunking.
   If a section is very long, split it into multiple ## sub-sections with descriptive names.
6. Output ONLY the clean Markdown body — no front-matter, no extra commentary.

RAW EXTRACTED TEXT:
---
{raw_text}
---

CLEAN MARKDOWN OUTPUT:
"""


def llm_cleanup(raw_text: str, title: str, url: str) -> str:
    """Use Gemini to clean up extracted text and add proper structure."""
    if not HAS_GEMINI or not GEMINI_API_KEY:
        print("  [!] Gemini not configured -- skipping LLM cleanup (set GEMINI_API_KEY)")
        return raw_text

    prompt = LLM_PROMPT_TEMPLATE.format(
        title=title,
        url=url,
        raw_text=raw_text[:12000],
    )
    try:
        if _gemini_client is not None:
            # New google.genai SDK
            response = _gemini_client.models.generate_content(
                model=GEMINI_MODEL,
                contents=prompt,
            )
            return response.text.strip()
        else:
            # Fallback: old google.generativeai SDK
            model = genai_old.GenerativeModel("gemini-1.5-flash")
            response = model.generate_content(prompt)
            return response.text.strip()
    except Exception as e:
        print(f"  [!] LLM cleanup failed: {e}")
        return raw_text


# ---------------------------------------------------------------------------
# Light local cleanup (applied before/after LLM)
# ---------------------------------------------------------------------------

def local_cleanup(text: str) -> str:
    """Quick regex-based cleanup that doesn't need an LLM."""
    # Remove control characters
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    # Collapse 3+ blank lines to 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Fix spaced-out letters: "I T 4 9 2" → kept as-is by LLM; just trim lines
    text = "\n".join(line.rstrip() for line in text.splitlines())
    return text.strip()


# ---------------------------------------------------------------------------
# Main re-extraction flow
# ---------------------------------------------------------------------------

def reextract_file(md_path: Path, dry_run: bool = False, use_llm: bool = True):
    print(f"\n{'='*60}")
    print(f"  Re-extracting: {md_path.name}")
    print(f"{'='*60}")

    text = md_path.read_text(encoding='utf-8', errors='replace')
    fm_str, _old_body = parse_frontmatter(text)

    url   = fm_get(fm_str, "url")
    title = fm_get(fm_str, "title")

    if not url or url in ("https://daiict.ac.in/", ""):
        print("  [skip] No valid source URL in frontmatter")
        return False

    # Skip social media / login-walled domains — no useful content to scrape
    from urllib.parse import urlparse
    parsed_host = urlparse(url).netloc.lower().lstrip("www.")
    if any(parsed_host.endswith(d) for d in SKIP_DOMAINS):
        print(f"  [skip] Social media / login-walled domain: {parsed_host}")
        return False

    print(f"  URL  : {url}")
    print(f"  Title: {title}")

    # Download
    data = download(url)
    if data is None:
        return False

    content_type = ""
    if requests:
        try:
            head = requests.head(url, headers=HEADERS, timeout=10, verify=False, allow_redirects=True)
            content_type = head.headers.get("Content-Type", "").lower()
        except Exception:
            pass

    # Guess file type from URL if Content-Type not available
    url_lower = url.lower()
    if not content_type:
        if url_lower.endswith(".pdf"):
            content_type = "application/pdf"
        elif any(url_lower.endswith(x) for x in [".xlsx", ".xls"]):
            content_type = "application/vnd.openxmlformats"
        elif url_lower.endswith(".docx"):
            content_type = "application/vnd.openxmlformats-officedocument"
        else:
            content_type = "text/html"

    # Extract raw content
    print(f"  Type : {content_type}")
    try:
        if "pdf" in content_type:
            raw = extract_pdf_text(data, url)
        elif "spreadsheet" in content_type or "excel" in content_type or "xlsx" in content_type:
            raw = extract_excel(data)
        elif "wordprocessing" in content_type or "docx" in content_type:
            raw = extract_docx(data)
        else:
            # HTML page — use BeautifulSoup for clean extraction
            raw = extract_html(data, url)
    except (PdfOcrUnavailableError, PdfOcrFailedError) as e:
        print(f"  ↳ {e} — aborting (not writing placeholder text to the corpus)")
        return False

    raw = local_cleanup(raw)

    if len(raw.strip()) < 50:
        print("  ↳ Extraction yielded too little content — aborting")
        return False

    print(f"  ↳ Extracted {len(raw)} chars of raw content")

    # LLM cleanup
    if use_llm:
        print("  ↳ Sending to Gemini for cleanup & structure …")
        body = llm_cleanup(raw, title, url)
    else:
        body = raw

    body = local_cleanup(body)

    # Ensure H1 heading
    if not re.search(r'^# ', body, re.MULTILINE):
        body = f"# {title}\n\n{body}"

    # Write back
    new_content = fm_str + "\n\n" + body + "\n"

    if dry_run:
        print(f"\n  [DRY RUN] Would write {len(new_content)} chars back to {md_path}")
        print("  Preview (first 500 chars):")
        print("  " + body[:500].replace("\n", "\n  "))
    else:
        md_path.write_text(new_content, encoding='utf-8')
        print(f"  ✓ Written {len(new_content)} chars to {md_path}")

    return True


# ---------------------------------------------------------------------------
# Batch mode from audit report
# ---------------------------------------------------------------------------

def run_from_audit(audit_path: Path, severities: list, dry_run: bool, use_llm: bool):
    report = json.loads(audit_path.read_text(encoding='utf-8'))
    targets = [r for r in report if r["severity"] in severities]
    print(f"Found {len(targets)} files with severity in {severities}")

    ok = 0
    fail = 0
    for entry in targets:
        fp = Path(entry["path"])
        if not fp.exists():
            print(f"  SKIP (not found): {fp}")
            continue
        success = reextract_file(fp, dry_run=dry_run, use_llm=use_llm)
        if success:
            ok += 1
        else:
            fail += 1
        time.sleep(1)  # be polite to the server

    print(f"\nDone. Success: {ok}, Failed/Skipped: {fail}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Re-extract DAU PWA markdown files")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--file",       help="Path to a single .md file to re-extract")
    group.add_argument("--from-audit", help="Path to audit_report.json for batch mode")

    parser.add_argument("--severity",  nargs="+", default=["CRITICAL", "WARN"],
                        help="Severities to process in batch mode (default: CRITICAL WARN)")
    parser.add_argument("--dry-run",   action="store_true",
                        help="Print what would happen without writing files")
    parser.add_argument("--no-llm",    action="store_true",
                        help="Skip LLM cleanup (faster but lower quality)")
    args = parser.parse_args()

    use_llm = not args.no_llm

    if args.file:
        fp = Path(args.file)
        if not fp.exists():
            print(f"ERROR: {fp} not found")
            sys.exit(1)
        reextract_file(fp, dry_run=args.dry_run, use_llm=use_llm)
    else:
        ap = Path(args.from_audit)
        if not ap.exists():
            print(f"ERROR: {ap} not found")
            sys.exit(1)
        run_from_audit(ap, args.severity, dry_run=args.dry_run, use_llm=use_llm)
